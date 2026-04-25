from docx import Document
from langchain_core.documents import Document as LangDocument

def load_docx(file):
    doc = Document(file)
    text = "\n".join([p.text for p in doc.paragraphs])

    return [LangDocument(page_content=text, metadata={"source": file.name})]